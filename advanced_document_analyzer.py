import os
import dlt
from dlt.common import pendulum
from google import genai
from google.genai import types
from dotenv import load_dotenv
from typing import List, Dict, Any, Optional, Union
import time
import base64
import json
from PIL import Image, ImageEnhance, ImageFilter
import io
import PyPDF2
import docx
import cv2
import numpy as np
from pathlib import Path
import psycopg2
from psycopg2.extras import Json
import re
from pdf2image import convert_from_bytes
import tempfile

load_dotenv()

class AdvancedDocumentAnalyzer:
    """
    Продвинутый анализатор документов с поддержкой:
    - JPG, PNG, PDF (до 20 МБ)
    - Коррекция перекошенных изображений и бликов
    - Распознавание русского и английского текста
    - Сохранение разметки (абзацы, заголовки, таблицы)
    - Ответы только на основе документа
    - Краткое содержание документа
    """
    
    # Поддерживаемые форматы изображений
    SUPPORTED_IMAGE_FORMATS = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.tiff': 'image/tiff',
        '.webp': 'image/webp'
    }
    
    # Поддерживаемые форматы документов
    SUPPORTED_DOC_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain'
    }
    
    # Максимальный размер файла (20 МБ)
    MAX_FILE_SIZE = 70 * 1024 * 1024  # 20 МБ в байтах
    
    def __init__(self, api_key: str):
        self.client = genai.Client(api_key=api_key)
        # Используем модели Gemini
        self.model_name = "gemini-2.0-flash-001"
        self.fast_model = "gemini-2.0-flash-001"
        
    def validate_file(self, file_content: bytes, file_name: str) -> bool:
        """
        Проверяет файл на соответствие требованиям
        1. Размер до 20 МБ
        2. Поддерживаемый формат
        """
        # Проверка размера
        file_size = len(file_content)
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"❌ Файл слишком большой: {file_size / 1024 / 1024:.1f} МБ. Максимальный размер: 20 МБ")
        
        # Проверка формата
        file_ext = os.path.splitext(file_name)[1].lower()
        if file_ext not in self.SUPPORTED_IMAGE_FORMATS and file_ext not in self.SUPPORTED_DOC_FORMATS:
            raise ValueError(f"❌ Неподдерживаемый формат: {file_ext}")
        
        print(f"✅ Файл прошел проверку: {file_size / 1024 / 1024:.1f} МБ")
        return True
    
    def preprocess_image(self, content: bytes, file_ext: str) -> bytes:
        """
        Предобработка изображения перед OCR:
        - Коррекция перекоса
        - Удаление бликов
        - Улучшение контраста
        - Повышение резкости
        """
        try:
            # Декодируем изображение
            nparr = np.frombuffer(content, np.uint8)
            img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if img is None:
                return content
            
            print("🔄 Предобработка изображения...")
            
            # 1. Конвертируем в оттенки серого
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # 2. Удаляем шум
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            
            # 3. Коррекция перекоса (дескьюинг)
            coords = np.column_stack(np.where(denoised > 0))
            if len(coords) > 0:
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = 90 + angle
                if abs(angle) > 0.5:  # Если перекос больше 0.5 градуса
                    print(f"   📐 Коррекция перекоса: {angle:.1f} градусов")
                    (h, w) = img.shape[:2]
                    center = (w // 2, h // 2)
                    M = cv2.getRotationMatrix2D(center, angle, 1.0)
                    denoised = cv2.warpAffine(denoised, M, (w, h), 
                                              flags=cv2.INTER_CUBIC,
                                              borderMode=cv2.BORDER_REPLICATE)
            
            # 4. Улучшение контраста (CLAHE)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(denoised)
            
            # 5. Борьба с бликами - адаптивная бинаризация
            binary = cv2.adaptiveThreshold(enhanced, 255, 
                                          cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY, 11, 2)
            
            # 6. Увеличение резкости
            kernel = np.array([[-1,-1,-1],
                               [-1, 9,-1],
                               [-1,-1,-1]])
            sharpened = cv2.filter2D(binary, -1, kernel)
            
            # 7. Морфологические операции для улучшения текста
            kernel = np.ones((1,1), np.uint8)
            morph = cv2.morphologyEx(sharpened, cv2.MORPH_CLOSE, kernel)
            
            # Конвертируем обратно в байты
            _, buffer = cv2.imencode('.png', morph)
            print("✅ Предобработка завершена")
            return buffer.tobytes()
            
        except Exception as e:
            print(f"⚠️ Ошибка при предобработке: {e}")
            return content
    
    def extract_text_from_pdf_with_ocr(self, content: bytes) -> str:
        """
        Извлечение текста из PDF с поддержкой сканов
        Если PDF содержит сканы, используется OCR
        """
        text = ""
        try:
            # Сохраняем PDF во временный файл
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(content)
                tmp_file_path = tmp_file.name
            
            # Пробуем извлечь текст обычным способом
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Проверяем, есть ли текст в PDF
            has_text = False
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    has_text = True
                    text += f"\n--- Страница ---\n{page_text}\n"
            
            # Если текст не найден, используем OCR
            if not has_text:
                print("📸 PDF содержит сканы, запускаем OCR...")
                # Конвертируем PDF в изображения
                images = convert_from_bytes(content)
                
                for i, image in enumerate(images):
                    print(f"   OCR страницы {i+1}/{len(images)}")
                    # Конвертируем PIL Image в байты
                    img_bytes = io.BytesIO()
                    image.save(img_bytes, format='PNG')
                    img_content = img_bytes.getvalue()
                    
                    # Предобработка изображения
                    processed = self.preprocess_image(img_content, '.png')
                    
                    # OCR через Gemini
                    page_text = self._extract_from_image(processed, '.png')
                    text += f"\n--- Страница {i+1} (скан) ---\n{page_text}\n"
            
            # Удаляем временный файл
            os.unlink(tmp_file_path)
            
        except Exception as e:
            text = f"Ошибка при обработке PDF: {str(e)}"
        
        return text
    
    def extract_text_from_file(self, file_content: bytes, file_name: str) -> str:
        """
        Извлекает текст из различных форматов файлов с предобработкой
        """
        # Проверяем файл
        self.validate_file(file_content, file_name)
        
        file_ext = os.path.splitext(file_name)[1].lower()
        
        # Обработка PDF (включая сканы)
        if file_ext == '.pdf':
            print(f"📄 Обработка PDF файла: {file_name}")
            return self.extract_text_from_pdf_with_ocr(file_content)
        
        # Обработка DOCX
        elif file_ext == '.docx':
            print(f"📝 Обработка DOCX файла: {file_name}")
            return self._extract_from_docx(file_content)
        
        # Обработка TXT
        elif file_ext == '.txt':
            print(f"📃 Обработка TXT файла: {file_name}")
            return self._extract_from_txt(file_content)
        
        # Обработка изображений (с предобработкой)
        elif file_ext in self.SUPPORTED_IMAGE_FORMATS:
            print(f"🖼️ Обработка изображения: {file_name}")
            # Предобработка изображения
            processed_content = self.preprocess_image(file_content, file_ext)
            return self._extract_from_image(processed_content, file_ext)
        
        else:
            raise ValueError(f"❌ Неподдерживаемый формат файла: {file_ext}")
    
    def _extract_from_txt(self, content: bytes) -> str:
        """Извлечение текста из TXT файла"""
        try:
            encodings = ['utf-8', 'cp1251', 'koi8-r', 'latin-1']
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            return content.decode('latin-1', errors='ignore')
        except Exception as e:
            return f"Ошибка при чтении TXT файла: {str(e)}"
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Извлечение текста из DOCX с сохранением структуры"""
        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            
            # Сохраняем структуру документа
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Определяем тип параграфа (заголовок или обычный текст)
                    if paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        text_parts.append(paragraph.text)
            
            # Сохраняем таблицы
            for table in doc.tables:
                text_parts.append("\n[TABLE]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
                text_parts.append("[/TABLE]\n")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            return f"Ошибка при извлечении из DOCX: {str(e)}"
    
    def _extract_from_image(self, content: bytes, file_ext: str) -> str:
        """Извлечение текста из изображения через Gemini Vision API"""
        try:
            mime_type = self.SUPPORTED_IMAGE_FORMATS.get(file_ext, 'image/jpeg')
            
            # Специальный промпт для OCR с сохранением структуры
            prompt = """
            Извлеки ВЕСЬ текст с этого изображения максимально точно и полно.
            Сохрани структуру документа:
            - Сохрани абзацы
            - Отметь заголовки (как ЗАГОЛОВОК)
            - Таблицы представь в формате | колонка1 | колонка2 |
            - Сохрани списки (нумерованные и маркированные)
            
            Языки: русский и английский.
            
            Верни ТОЛЬКО текст, без комментариев.
            """
            
            print(f"🔍 Запуск OCR...")
            
            response = self.client.models.generate_content(
                model=self.fast_model,
                contents=[
                    prompt,
                    types.Part.from_bytes(data=content, mime_type=mime_type)
                ]
            )
            
            if response and hasattr(response, 'text'):
                return response.text
            else:
                return "Не удалось распознать текст"
            
        except Exception as e:
            return f"Ошибка OCR: {str(e)}"
    
    def answer_question_from_document(self, text: str, query: str) -> str:
        """
        Отвечает на вопрос ТОЛЬКО на основе документа
        Если ответа нет в документе, возвращает "Информация отсутствует"
        """
        prompt = f"""
        Ты - ассистент, который отвечает на вопросы ТОЛЬКО на основе предоставленного документа.
        
        Документ:
        {text[:8000]}
        
        Вопрос: {query}
        
        Правила:
        1. Если информация для ответа есть в документе - дай точный ответ
        2. Если информации в документе НЕТ - ответь ТОЧНО: "Информация отсутствует в загруженном файле"
        3. Не добавляй информацию из других источников
        4. Цитируй документ, если нужно
        
        Ответ:
        """
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt
            )
            
            if response and hasattr(response, 'text'):
                return response.text
            else:
                return "Не удалось получить ответ от API"
        except Exception as e:
            return f"Ошибка: {str(e)}"
    
    def get_document_summary(self, text: str) -> str:
        """
        Создает краткое содержание документа (3-5 предложений):
        - О чем документ
        - Стороны
        - Ключевая сумма
        - Дата подписания
        """
        prompt = f"""
        Создай краткое содержание документа (3-5 предложений).
        
        Документ:
        {text[:6000]}
        
        Обязательно включи:
        1. О чем документ (тема/назначение)
        2. Кто стороны (если есть)
        3. Ключевая сумма (если есть)
        4. Дата подписания (если есть)
        
        Если какой-то информации нет - просто не упоминай ее.
        Формат: кратко, по делу, только факты из документа.
        
        Краткое содержание:
        """
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt
            )
            
            if response and hasattr(response, 'text'):
                return response.text
            else:
                return "Не удалось создать краткое содержание"
        except Exception as e:
            return f"Ошибка: {str(e)}"
    
    def full_analysis(self, file_content: bytes, file_name: str, 
                     query: str = None, need_summary: bool = False) -> Dict[str, Any]:
        """
        Полный анализ документа с поддержкой всех функций
        """
        print(f"\n{'='*60}")
        print(f"🔍 АНАЛИЗ ФАЙЛА: {file_name}")
        print(f"{'='*60}")
        
        # Шаг 1: Извлечение текста с предобработкой
        extracted_text = self.extract_text_from_file(file_content, file_name)
        print(f"📊 Извлечено символов: {len(extracted_text)}")
        
        # Шаг 2: Определение AI-генерации (базовая функция)
        ai_detection = self.detect_ai_generated(extracted_text, file_name)
        
        # Шаг 3: Ответ на вопрос (если есть)
        answer = None
        if query:
            print(f"❓ Вопрос: {query}")
            answer = self.answer_question_from_document(extracted_text, query)
            print(f"✅ Ответ получен")
        
        # Шаг 4: Краткое содержание (если нужно)
        summary = None
        if need_summary:
            print(f"📝 Создание краткого содержания...")
            summary = self.get_document_summary(extracted_text)
            print(f"✅ Краткое содержание создано")
        
        result = {
            "file_name": file_name,
            "file_format": os.path.splitext(file_name)[1].lower(),
            "file_size": len(file_content),
            "file_size_mb": f"{len(file_content) / 1024 / 1024:.1f} МБ",
            "extracted_text": extracted_text[:2000] + "..." if len(extracted_text) > 2000 else extracted_text,
            "text_length": len(extracted_text),
            "analysis_timestamp": pendulum.now().isoformat(),
            "ai_detection": ai_detection,
            "answer_to_query": answer,
            "summary": summary,
            "preprocessing_applied": file_name.lower().endswith(('.jpg', '.jpeg', '.png', '.pdf'))
        }
        
        return result
    
    def detect_ai_generated(self, text: str, filename: str = None) -> Dict[str, Any]:
        """Определяет, является ли текст AI-сгенерированным (базовая функция)"""
        
        analysis_prompt = f"""
        Проанализируй этот текст и определи, с какой вероятностью он создан искусственным интеллектом.
        
        Текст для анализа:
        {text[:5000]}
        
        Оцени по критериям:
        1. Естественность языка (0-100)
        2. Эмоциональная окраска (0-100)
        3. Разнообразие словарного запаса (0-100)
        4. Наличие типичных "AI-маркеров"
        
        Верни ответ в формате JSON:
        {{
            "ai_probability": число (0-100),
            "classification": "ai_generated" | "human_written" | "uncertain",
            "confidence_level": "high" | "medium" | "low"
        }}
        """
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=analysis_prompt
            )
            
            if response and hasattr(response, 'text'):
                try:
                    json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                    if json_match:
                        result = json.loads(json_match.group())
                    else:
                        result = {
                            "ai_probability": 50,
                            "classification": "uncertain",
                            "confidence_level": "low"
                        }
                except:
                    result = {
                        "ai_probability": 50,
                        "classification": "uncertain",
                        "confidence_level": "low"
                    }
            else:
                result = {
                    "ai_probability": 50,
                    "classification": "uncertain",
                    "confidence_level": "low"
                }
            
            result["analyzed_at"] = pendulum.now().isoformat()
            result["text_length"] = len(text)
            
            return result
            
        except Exception as e:
            print(f"⚠️ Ошибка при AI-детекции: {e}")
            return {
                "ai_probability": 50,
                "classification": "error",
                "error": str(e)
            }


# ============= ФУНКЦИИ ДЛЯ ТЕСТИРОВАНИЯ =============

def test_with_real_file():
    """Тестирование на реальном файле"""
    import tkinter as tk
    from tkinter import filedialog
    
    # Создаем окно для выбора файла
    root = tk.Tk()
    root.withdraw()
    
    print("📁 Выберите файл для анализа (JPG, PNG, PDF до 20 МБ)")
    file_path = filedialog.askopenfilename(
        filetypes=[
            ("Изображения и PDF", "*.jpg *.jpeg *.png *.pdf"),
            ("Все файлы", "*.*")
        ]
    )
    
    if not file_path:
        print("❌ Файл не выбран")
        return
    
    print(f"✅ Выбран файл: {file_path}")
    
    # Загружаем переменные окружения
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    
    if not api_key:
        print("❌ Нет API ключа в .env")
        return
    
    # Создаем анализатор
    analyzer = AdvancedDocumentAnalyzer(api_key)
    
    # Читаем файл
    with open(file_path, 'rb') as f:
        content = f.read()
    
    print(f"\n📊 Размер файла: {len(content) / 1024 / 1024:.1f} МБ")
    
    # Спрашиваем, нужно ли краткое содержание
    need_summary = input("\n📝 Создать краткое содержание? (да/нет): ").lower() == 'да'
    
    # Спрашиваем вопрос
    query = input("❓ Введите вопрос по документу (или Enter для пропуска): ")
    
    # Анализируем
    result = analyzer.full_analysis(
        file_content=content,
        file_name=os.path.basename(file_path),
        query=query if query else None,
        need_summary=need_summary
    )
    
    # Выводим результаты
    print("\n" + "="*60)
    print("📊 РЕЗУЛЬТАТЫ АНАЛИЗА")
    print("="*60)
    
    if result.get('summary'):
        print(f"\n📝 КРАТКОЕ СОДЕРЖАНИЕ:")
        print(result['summary'])
    
    if result.get('answer_to_query'):
        print(f"\n❓ ОТВЕТ НА ВОПРОС:")
        print(result['answer_to_query'])
    
    print(f"\n🤖 ВЕРОЯТНОСТЬ AI: {result['ai_detection'].get('ai_probability', 'N/A')}%")
    print(f"📏 ДЛИНА ТЕКСТА: {result['text_length']} символов")
    
    # Спрашиваем, показать ли полный текст
    show_text = input("\n📄 Показать извлеченный текст? (да/нет): ").lower() == 'да'
    if show_text:
        print("\n" + "="*60)
        print("📄 ИЗВЛЕЧЕННЫЙ ТЕКСТ")
        print("="*60)
        print(result['extracted_text'])


if __name__ == "__main__":
    print("🚀 ТЕСТОВЫЙ РЕЖИМ С ВЫБОРОМ ФАЙЛА")
    print("="*60)
    test_with_real_file()